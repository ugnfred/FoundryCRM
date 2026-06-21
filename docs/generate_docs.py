"""
Converts docs/*.md files to formatted Word (.docx) documents.
Run from project root:  python docs/generate_docs.py
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand colours ──────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x1E, 0x40, 0xAF)   # indigo-800
HEADER_BG = RGBColor(0x1E, 0x40, 0xAF)
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW   = RGBColor(0xEF, 0xF6, 0xFF)   # blue-50
BORDER    = RGBColor(0xCB, 0xD5, 0xE1)   # slate-300


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), 'CBD5E1')
        borders.append(el)
    tcPr.append(borders)


def apply_table_style(table):
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell)
            if i == 0:
                set_cell_bg(cell, HEADER_BG)
            elif i % 2 == 0:
                set_cell_bg(cell, ALT_ROW)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    if i == 0:
                        run.font.color.rgb = HEADER_FG
                        run.font.bold      = True
                    run.font.size = Pt(9)


def add_cover_page(doc, title, subtitle="Foundry ERP — Royal Met Alloys"):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph(subtitle)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.runs[0]
    r.font.size  = Pt(12)
    r.font.color.rgb = PRIMARY

    h = doc.add_paragraph(title)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.runs[0]
    r.font.size  = Pt(28)
    r.font.bold  = True
    r.font.color.rgb = PRIMARY

    doc.add_paragraph()
    ver = doc.add_paragraph("Version 1.0  ·  June 2026")
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.runs[0].font.size = Pt(11)
    ver.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()


def parse_and_add(doc, md_text):
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Page break on ---  (only if it's a section separator, not table border)
        if re.match(r'^---+$', line.strip()) and i > 0:
            # skip — just a visual divider in MD
            i += 1
            continue

        # ── Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # Strip anchor links like {#anchor}
            text  = re.sub(r'\s*\{#[^}]+\}', '', text)
            style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}.get(level, 'Heading 4')
            p = doc.add_heading(text, level=level)
            p.style.font.color.rgb = PRIMARY
            i += 1
            continue

        # ── Tables (GitHub-flavored: lines with |)
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Remove separator row (|---|---|)
            rows = [r for r in table_lines if not re.match(r'^\|[\s\-\|:]+\|$', r.strip())]
            if not rows:
                continue
            parsed = []
            for r in rows:
                cells = [c.strip() for c in r.strip().strip('|').split('|')]
                parsed.append(cells)
            col_count = max(len(r) for r in parsed)
            table = doc.add_table(rows=len(parsed), cols=col_count)
            table.style = 'Table Grid'
            for ri, row_data in enumerate(parsed):
                for ci, cell_text in enumerate(row_data):
                    if ci < col_count:
                        # Strip markdown bold/code/emoji from cell text
                        cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                        table.rows[ri].cells[ci].text = cell_text
            apply_table_style(table)
            doc.add_paragraph()
            continue

        # ── Code block
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph('\n'.join(code_lines))
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
            # Grey background via shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'F1F5F9')
            pPr.append(shd)
            doc.add_paragraph()
            continue

        # ── Bullet list (- item or * item)
        m = re.match(r'^(\s*)[*\-]\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            text   = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text   = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5 * (indent + 1))
            run = p.add_run(text)
            run.font.size = Pt(10)
            i += 1
            continue

        # ── Numbered list (1. item)
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            text = m.group(1)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            run.font.size = Pt(10)
            i += 1
            continue

        # ── Blockquote (> text)
        m = re.match(r'^>\s+(.*)', line)
        if m:
            text = m.group(1)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x47, 0x56, 0x69)
                run.font.size = Pt(10)
            i += 1
            continue

        # ── Blank line
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph (skip the H1 title at top — covered by cover page)
        if re.match(r'^#\s', line):
            i += 1
            continue

        # ── Regular text
        text = line
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'`(.*?)`',       r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # links → text
        if text.strip():
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(10)
        i += 1


def generate(md_path, docx_path, title):
    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default body font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    for level in range(1, 5):
        style_name = f'Heading {level}'
        s = doc.styles[style_name]
        s.font.name      = 'Calibri'
        s.font.color.rgb = PRIMARY
        s.font.bold      = True

    add_cover_page(doc, title)
    parse_and_add(doc, content)

    doc.save(docx_path)
    print(f"  OK  {docx_path}")


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))

    files = [
        ('ROLE_MATRIX.md',  'ROLE_MATRIX.docx',  'Role & Access Matrix'),
        ('USER_MANUAL.md',  'USER_MANUAL.docx',  'User Manual'),
    ]

    print("Generating Word documents...")
    for md_name, docx_name, title in files:
        generate(
            os.path.join(base, md_name),
            os.path.join(base, docx_name),
            title,
        )
    print("Done. Word files saved in docs/")
